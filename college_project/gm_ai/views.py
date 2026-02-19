from django.contrib.auth import logout
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.hashers import make_password
from django.contrib.auth import authenticate, login
from .models import CustomUser
from django.contrib.auth.decorators import login_required
from django.conf import settings
from django.http import FileResponse
from django.utils import timezone
from docx import Document
from datetime import datetime
import os
from .models import Contract

def home(request):
    return render(request, 'home.html')

# REGISTER
def register_view(request):
    if request.method == "POST":

        username = request.POST.get("login")
        full_name = request.POST.get("full_name")
        phone = request.POST.get("phone")
        passport_id = request.POST.get("passport_id")
        permanent_address = request.POST.get("permanent_address")
        password = request.POST.get("password")

        if CustomUser.objects.filter(username=username).exists():
            messages.error(request, "Bu login band!")
            return redirect("register")

        CustomUser.objects.create(
            username=username,
            full_name=full_name,
            phone=phone,
            passport_id=passport_id,
            permanent_address=permanent_address,
            password=make_password(password)
        )

        messages.success(request, "Ro'yxatdan o'tdingiz! Login qiling.")
        return redirect("login")

    return render(request, "register.html")


# LOGIN
def login_view(request):
    if request.method == "POST":

        username = request.POST.get("login")
        password = request.POST.get("password")

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            return redirect("home")
        else:
            messages.error(request, "Login yoki parol noto‘g‘ri!")

    return render(request, "login.html")


# LOGOUT
def logout_view(request):
    logout(request)
    return redirect("home")

#all cars
def cars_list(request):
    return render(request, 'cars_list.html')

#Captiva
def captiva_detail(request):
    return render(request, 'models/captiva.html')

def traverse_detail(request):
    return render(request, 'models/traverse.html')

def tracker_detail(request):
    return render(request, 'models/tracker.html')

def onix_detail(request):
    return render(request, 'models/onix.html')
def tahoe_detail(request):
    return render(request, 'models/tahoe.html')

from .models import Autosalon

def captiva_config(request):
    autosalons = Autosalon.objects.all()
    return render(request, 'models/captiva_config.html', {
        "autosalons": autosalons
    })
def onix_config(request):
    autosalons = Autosalon.objects.all()
    return render(request, 'models/onix_config.html', {
        "autosalons": autosalons
    })
def tracker_config(request):
    autosalons = Autosalon.objects.all()
    return render(request, 'models/tracker_config.html', {
        "autosalons": autosalons
    })
def tahoe_config(request):
    autosalons = Autosalon.objects.all()
    return render(request, 'models/tahoe_config.html', {
        "autosalons": autosalons
    })
def labo_config(request):
    autosalons = Autosalon.objects.all()
    return render(request, 'models/labo_config.html', {
        "autosalons": autosalons
    })
def damas_config(request):
    autosalons = Autosalon.objects.all()
    return render(request, 'models/damas_config.html', {
        "autosalons": autosalons
    })



#buy view
@login_required
def buy_car(request):
    if request.method == "POST":

        user = request.user
        model = request.POST.get("model")
        dealer_id = request.POST.get("dealer")
        color = request.POST.get("color")
        modification = request.POST.get("modification")

        # ------------------------
        # Autosalonni olish
        # ------------------------
        try:
            dealer = Autosalon.objects.get(id=dealer_id)
        except Autosalon.DoesNotExist:
            return redirect("home")

        # ------------------------
        # Contract Number Generatsiya
        # Format: UZM-2026-0001
        # ------------------------
        year = timezone.now().year
        last_contract = Contract.objects.filter(
            created_at__year=year
        ).order_by("-id").first()

        if last_contract:
            last_number = int(last_contract.contract_number.split("-")[-1])
            new_number = last_number + 1
        else:
            new_number = 1

        contract_number = f"UZM-{year}-{str(new_number).zfill(4)}"

        # ------------------------
        # Contract DB ga yozish
        # ------------------------
        contract = Contract.objects.create(
            user=user,
            autosalon=dealer,
            model=model,
            modification=modification,
            color=color,
            contract_number=contract_number
        )

        # ------------------------
        # DOCX Generatsiya
        # ------------------------
        document = Document()

        document.add_heading('AVTOMOBIL SOTIB OLISH SHARTNOMASI', level=1)
        document.add_paragraph(f"Shartnoma raqami: {contract.contract_number}")
        document.add_paragraph(f"Sana: {datetime.now().strftime('%d.%m.%Y')}")
        document.add_paragraph("")

        document.add_paragraph(f"Xaridor: {user.full_name}")
        document.add_paragraph(f"Passport ID: {user.passport_id}")
        document.add_paragraph(f"Telefon: {user.phone}")
        document.add_paragraph(f"Manzil: {user.permanent_address}")
        document.add_paragraph("")

        document.add_paragraph(f"Avtomobil: Chevrolet {model}")
        document.add_paragraph(f"Modifikatsiya: {modification.upper()}")
        document.add_paragraph(f"Rangi: {color}")
        document.add_paragraph("")

        document.add_paragraph(f"Avtosalon: {dealer.name}")
        document.add_paragraph(f"Shahar: {dealer.city}")
        document.add_paragraph(f"Manzil: {dealer.address}")
        document.add_paragraph("")

        document.add_paragraph("Tomonlar yuqoridagi ma’lumotlarga asosan shartnoma tuzdilar.")

        # ------------------------
        # Papka yaratish
        # ------------------------
        contracts_path = os.path.join(settings.MEDIA_ROOT, "contracts")
        os.makedirs(contracts_path, exist_ok=True)

        filename = f"{contract.contract_number}.docx"
        filepath = os.path.join(contracts_path, filename)

        document.save(filepath)

        return FileResponse(open(filepath, 'rb'), as_attachment=True)

    return redirect("home")

import json
import google.generativeai as genai
from django.http import JsonResponse
from django.conf import settings

# Configure Gemini
genai.configure(api_key=settings.GEMINI_API_KEY)

model = genai.GenerativeModel("gemini-3-flash-preview")


def chat(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    try:
        data = json.loads(request.body)
        message = data.get("message", "").strip()

        if not message:
            return JsonResponse({"response": "Iltimos, savol yozing."})

        prompt = f"""
        Sen UzAuto rasmiy AI yordamchisisan.
        Professional o‘zbek tilida javob ber.
        captiva narxi: LT2 pozitsiyaniki 304,900,000
                       PRIMIER pozitsiyaniki 334,900,000
        onix narxi:LS MT — 161,900,000 so‘m 
                   3LT MT — 169,750,000 so‘m
                   LTZ TURBO AT — 184,899,840 so‘m 
                   PREMIER 2 TURBO AT PLUS — 206,640,160 so‘m 
        tracker narxi: TRK LS PLUS — 215,951,360 so‘m
                       TRK LTZ PLUS — 229,108,480 so‘m
                       TRK PREMIER PLUS — 252,656,160 so‘m
                       TRK REDLINE — 260,474,080 so‘m
        tahoe narxi : 1,464,500,000 so`m
        Savol: {message}
        """

        response = model.generate_content(prompt)

        # SAFER extraction
        if response.candidates:
            reply = response.candidates[0].content.parts[0].text
        else:
            reply = "Javob topilmadi."

        return JsonResponse({"response": reply})

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
